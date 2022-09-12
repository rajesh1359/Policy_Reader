from docx import Document
import pdfplumber
document = r"E:\NewDownloads\DemoDownloads\NEW PDF\bajaj\New folder\4w comm.docx"

#    data = [[" ".join(cell.text.split()) for cell in row.cells] for row in table.rows]
Acronym_Map_BJ = {'BAJAJ': 'BAJAJ', 'TW': 'TWO WHEELER', 'PC': 'Private Car', 'COMM': "Passenger Carrying", 'TP': 'TP',
                  'PKG': 'Package'}

INSURANCE_TYPE_OLD = {
    'BAJAJ_TW': ['transcript', 'two', 'wheeler'],
    'BAJAJ_PC': ['transcript', 'private', 'car'],
    'BAJAJ_COMM': ['transcript', 'commercial', 'vehicle', 'policy']
}
INSURANCE_TYPE_NEW = {
    'BAJAJ_TW': ['two', 'wheeler'],
    'BAJAJ_PC': ['private', 'car'],
    'BAJAJ_COMM': ['commercial', 'vehicle', 'policy']
}

POLICY_MAP = {
    'OLD': INSURANCE_TYPE_OLD,
    'NEW': INSURANCE_TYPE_NEW
}

OLD_VS_NEW = {
    'TW': ['TP'],
    'PC': ['TP'],
    'COMM': ['TP']
}
 #   print(data)
_DEBUG= False
import re
class BajajDataExtractor:
    def __init__(self,file_path):
        self.file_path = file_path
        self.word_doc = Document(self.file_path)
        self.policy_version ='NEW'
        self.raw_data = pdfplumber.open(self.file_path[:-5]+".pdf")
    def get_policy_type(self):
        for pr in self.word_doc.paragraphs:
            row= pr.text.lower()
            if 'policy schedule cum' in row:
                for policy_type, attrib in POLICY_MAP[self.policy_version].items():
                    check_list = set([attr in row for attr in attrib])
                    if len(check_list) == 1 and list(check_list)[0]:
                        if 'liability' in row:
                            policy_type = policy_type + "_TP"
                        else:
                            policy_type = policy_type + "_PKG"
                        # print('POLICY TYPE ',policy_type)
                        if policy_type.split('_')[1]=='COMM':
                            insured_table = self.get_a_table(
                                ['INSURED DETAILS', 'POLICY DETAILS', 'Period of Insurance', 'Invoice Number'])
                            if insured_table is not None:
                                m=self.get_row_value('Vehicle Type', insured_table)
                                if len(m)!=0:
                                    use_string =m[0]
                                    if 'passenger' in use_string.lower():
                                        vtype= 'Passanger Carrying'
                                    else:
                                        if 'pubic' in use_string.lower():
                                            vtype = 'Goods - Public Carrier'
                                        else:
                                            vtype = 'Goods - PVT Carrier'

                                    return Acronym_Map_BJ[policy_type.split('_')[0]],vtype,Acronym_Map_BJ[policy_type.split('_')[2]]
                        return Acronym_Map_BJ[policy_type.split('_')[0]], Acronym_Map_BJ[policy_type.split('_')[1]], \
                               Acronym_Map_BJ[policy_type.split('_')[2]]
        return None,None,None
    def get_a_table(self,identifiers):

        for table in self.word_doc.tables:
            temp_identifier = identifiers
            for row in table.rows:
                for cell in row.cells:

                    for keyword in identifiers:
                        if keyword in " ".join(cell.text.split()):
                            temp_identifier.remove(keyword)
            if len(temp_identifier)==0:
                return table
            if _DEBUG:
                print(temp_identifier)
        return None

    def table_data(self, table):
        data = [[cell.text for cell in row.cells] for row in table.rows]
        return data
    def get_next_unique_cell(self,cellnum,row):
        row_cells = row.cells
        row_cells_count = len(row_cells)
        found=False
        cell_val = row_cells[cellnum].text
        for clm in range(cellnum+1,row_cells_count):
            if row_cells[clm].text != cell_val:
                found = True
                return row_cells[clm].text


    def get_row_value(self, key, tb, cell_type = 'Row', keep_going=False,regex=False,reg_value="(.*)"):
        #if keep_going:
        values=[]
        if cell_type == 'Row':
            tc_cell_line = tb.rows
        else:
            tc_cell_line=tb.columns
        cell_considered = []
        for row in tc_cell_line:
            cell_num = 0

            for cell in row.cells:
                cell_text= " ".join(cell.text.split())
                if key in  cell_text and not cell_text in cell_considered:
                    value = self.get_next_unique_cell(cell_num,row)
                    if value is not None:
                    #value = row.cells[cell_num+1].text
                        if regex:
                            m = re.match(reg_value,value)
                            if m is not None:
                                values.append(m.groups()[0])
                                if not keep_going:
                                    return values
                            else:
                                pass
                        else:

                            values.append(value.strip())
                            cell_considered.append(cell_text)
                            if not keep_going:
                                return values

                cell_num +=1

        if len(values)==0:
            m = self.get_value_from_row(tb,".*"+key+"\s*.*\s+([+-]?([0-9]*[.])?[0-9]+).*")
            if m is not None:
                print('Primium ',m.groups()[0])
        return values
    def get_value_from_row(self,tb,regex):

        for rw in tb.rows:
            for cell in rw.cells:
                str1=" ".join(cell.text.split())
                m = re.match(regex,str1.lower())
                if m is not None:
                    return m.groups()[0]
    def get_basic_info(self):
        insured_table = self.get_a_table(['Name', 'Mailing Address', 'Mobile Number', 'Email ID'])
        if insured_table is not None:
            table_data = self.table_data(insured_table)
            if _DEBUG:
                print('Insured MObile :', self.get_row_value('Mobile Number', insured_table))
            mob_num =None
            m = self.get_row_value('Mobile Number', insured_table)
            if len(m) != 0:
                mob_num = m[0].strip()
            return mob_num
    def get_insurance_details(self):
        insured_table = self.get_a_table(['INSURED DETAILS','POLICY DETAILS','Period of Insurance','Invoice Number'])
        if insured_table is not None:
            table_data = self.table_data(insured_table)
            if _DEBUG:
                print(table_data)
            # print('Insured Name :', self.get_row_value('Insured Name',insured_table))
            # print('Policy Issued on :', self.get_row_value('Policy Issued on', insured_table, regex=True,reg_value='.*(\d{2}-\d{2}-\d{4})\s+.*'))
            # print('Policy Number :', self.get_row_value('Policy Number', insured_table))
            # # print('Policy Issued on :', self.get_row_value('Policy Issued on', insured_table))
            # print('Policy Expired On ',self.get_value_from_row(insured_table,'.*(\d{2}-\d{2}-\d{4})\s*midnight.*'))
            name,issue_date,policy_number,expiry_date = None,None,None,None
            m = self.get_row_value('Insured Name',insured_table)
            if len(m)!=0:
                name = m[0].strip()
            m = self.get_row_value('Policy Issued on', insured_table, regex=True,reg_value='.*(\d{2}-\d{2}-\d{4})\s+.*')
            if len(m) != 0:
                issue_date = m[0].strip()
            m = self.get_row_value('Policy Number', insured_table)
            if len(m) != 0:
                policy_number = m[0].strip()
            m = self.get_value_from_row(insured_table,'.*(\d{2}-\d{2}-\d{4})\s*midnight.*')
            if m is not None:
                expiry_date = m.strip()
            return name,issue_date,policy_number,expiry_date
    def get_clean_text(self, txt):
        return txt.replace('\n', ' ').replace('- ', '')

    def get_next_not_none(self,remaining_row):
        for cell in remaining_row:
            if cell.strip() != 'NONE':
                return cell
        return None
    def get_value_from_price_table(self,table,key,repeat=False):
        TotalValue = None
        for row in table:

            for cell_index, cell in enumerate(row):
                m= re.match(key,str(cell.strip()))
                if m is not None:
                    value = self.get_next_not_none(row[cell_index+1:])
                    if (value is not None and 'NA' not in value) and len(value.strip())!=0:
                        if TotalValue is None:
                            TotalValue=0
                        TotalValue+=float(value.replace(',',''))
                        if repeat:
                            pass
                        else:
                            return TotalValue
        return TotalValue
    def get_price_table(self):
        customer_name = None
        price_Page = None
        for page in self.raw_data.pages:
            raw_text = page.extract_text(x_tolerance=1).upper()
            # raw_text_list = raw_text.lower().split("\n")
            if 'OWN DAMAGE' in raw_text and 'FINAL PREMIUM' in raw_text:
                price_Page = page
                break
        if price_Page is not None:
            sc_res = price_Page.search(".*(Own Damage).*", regex=True, case=False, x_tolerance=1)
            if len(sc_res) != 0:
                top_y = sc_res[0]['top'] - 2 + sc_res[0]['top'] - sc_res[0]['bottom']
                x0 = 10
                if top_y + 150 < price_Page.height:
                    bottom_y = top_y + 150
                else:
                    bottom_y = price_Page.height
                x1 = price_Page.width
                new_page = price_Page.crop((x0, top_y, x1, bottom_y))
                print("{0} {1} {2} {3}".format(x0,top_y, x1, bottom_y))
                tbls = new_page.extract_tables({'text_x_tolerance': 1})
                for tbl in tbls:
                    found_check_list = []
                    for row in tbl:
                        for cell in row:
                            if 'own damage' in str(cell).lower():
                                found_check_list.append(1)
                    if 1 in found_check_list:
                        clean_table = []
                        for row in tbl:
                            clean_row = [self.get_clean_text(str(data).upper()) for data in row]
                            clean_table.append(clean_row)
                        return clean_table
            return None
    def get_prices(self):
        pc_table = self.get_price_table()
        if pc_table is not None:
            '''
            print('OD PRIMUIM ',self.get_value_from_price_table(pc_table,'.*OWN DAMAGE PREMIUM:*(?!\()'))
            print('TOTAL GST ', self.get_value_from_price_table(pc_table,'.*GST.*',True))
            print('Final Premium ', self.get_value_from_price_table(pc_table,'.*FINAL\s*PREMIUM.*'))
            print('TP Premium ', self.get_value_from_price_table(pc_table,'.*BASIC\s*THIRD\s*PARTY\s*LIABILITY.*'))

            '''
            ODP = self.get_value_from_price_table(pc_table, '.*TOTAL OWN DAMAGE:*(?!\()')
            GST= self.get_value_from_price_table(pc_table, '.*GST.*', True)
            FP = self.get_value_from_price_table(pc_table, '.*FINAL\s*PREMIUM.*')
            TP = self.get_value_from_price_table(pc_table, '.*BASIC\s*THIRD\s*PARTY\s*LIABILITY.*')
            print(pc_table)
            if TP is None:
                TP=0
            if ODP is None:
                ODP=0
            print(FP)
            print(GST)
            NetPremium = FP - GST
            return FP,TP,ODP,NetPremium
        return None,None,None,None
    def get_premium_details(self):
        insured_table = self.get_a_table(['Total Own Damage', 'Basic Third Party Liability','Final Premium'])
        if insured_table is not None:
            table_data = self.table_data(insured_table)
            print(table_data)
            if _DEBUG:
                print(table_data)
                print('Total Own Damage :',self.get_row_value('Total Own Damage',insured_table))
                print('Basic Third Party Liability :',self.get_row_value('Basic Third Party Liability', insured_table))
                # print('Total Liability Premium :',self.get_row_value('Total Liability Premium', insured_table))
                # print('Total Premium :',self.get_row_value('Total Premium', insured_table))
                print('Final Premium :',self.get_row_value('Final Premium', insured_table))
                print('GST on Premium :',self.get_row_value('GST', insured_table,keep_going=True))
            TP, OD, FP, NET = None,None,None,None
            m = self.get_row_value('Basic Third Party Liability', insured_table)
            if len(m) != 0:
                TP = m[0].strip()
                if TP =='NA':
                    TP = '0'
            m = self.get_row_value('Total Own Damage', insured_table)
            if len(m) != 0:
                OD = m[0].strip()
                if OD =='NA':
                    OD = '0'
            m = self.get_row_value('Final Premium', insured_table)
            print(m)
            if len(m) != 0:
                FP = m[0].strip()
                if FP == 'NA':
                    FP = '0'

            GST=0
            m = self.get_row_value('GST', insured_table, keep_going=True)
            for gst in m:
                if len(gst.strip()) !=0 and gst.strip()!='NA':
                    GST+=float(gst.replace(',',''))

            if FP is None:
                FP = self.get_value_from_row(insured_table,".*final\s*premium\s*([+-]?([0-9]*[.])?[0-9]+).*")
            if GST is None:
                GST = self.get_value_from_row(insured_table,".*gst\s*([+-]?([0-9]*[.])?[0-9]+).*")
            if TP is None:
                TP = self.get_value_from_row(insured_table,".*basic\s*third\s*party\s*liability\s*([+-]?([0-9]*[.])?[0-9]+).*")
            if OD is None:
                OD = self.get_value_from_row(insured_table,".*own\s*damage\s*([+-]?([0-9]*[.])?[0-9]+).*")

            if FP is not None:
                FP = float(str(FP).replace(',', ''))
            if GST is not None:
                GST = float(str(GST).replace(',', ''))
            if TP is not None:
                TP = float(str(TP).replace(',', ''))
            if OD is not None:
                OD = float(str(OD).replace(',', ''))
            NET = FP - GST
            return FP, TP, OD, NET

    def get_agency_details(self):
        insured_table = self.get_a_table(['Agency Code', 'Phone Number','Sub IMD Code'])
        if insured_table is not None:
            table_data = self.table_data(insured_table)
            if _DEBUG:
                print(table_data)
            print('Sub IMD Code :',self.get_row_value('Sub IMD Code',insured_table))
            code =None
            m = self.get_row_value('Sub IMD Code',insured_table)
            if len(m)!=0:
                code = m[0]
            return code

    def get_bike_details(self):
        insured_table = self.get_a_table(['Vehicle Make', 'Vehicle Model', 'Registration'])
        if insured_table is not None:
            table_data = self.table_data(insured_table)
            if _DEBUG:
            #print(table_data)
                print('Vehicle Make:', self.get_row_value('Vehicle Make', insured_table, cell_type='Column'))
                print('Registration Number:', self.get_row_value('Registration Number', insured_table, cell_type='Column'))
                print('Vehicle Model:', self.get_row_value('Vehicle Model', insured_table, cell_type='Column'))
            make,model,reg_num=None,None,None
            m = self.get_row_value('Vehicle Make', insured_table, cell_type='Column')
            if len(m) != 0:
                make = m[0]

            m = self.get_row_value('Registration Number', insured_table, cell_type='Column')
            if len(m) != 0:
                reg_num = m[0]

            m = self.get_row_value('Vehicle Model', insured_table, cell_type='Column')
            if len(m) != 0:
                model = m[0]
            return make,model,reg_num
    def get_all_details(self):
        data_dict = {
            'A': None,
            'B': None,
            'C': None,
            'D': None,
            'E': None,
            'F': None,
            'G': None,
            'H': None,
            'I': None,
            'J': None,
            'K': None,
            'L': None,
            'M': None,
            'N': None,
            'O': None,
            'P': None,
            'Q': None,
            'R': None,
            'S': None,
            'T': None,
            'U': None,
            'V': None,
            'W': None,
            'X': None
        }

        data_dict['B'],data_dict['A'],data_dict['D'],data_dict['E'] = self.get_insurance_details()
        data_dict['Q'] = self.get_agency_details()
        data_dict['J'],data_dict['K'],data_dict['C'] = self.get_bike_details()
        data_dict['F'] = self.get_basic_info()
        data_dict['I'],data_dict['G'],data_dict['H'] = self.get_policy_type()
        # print(self.get_prices())
        data_dict['L'], data_dict['M'], data_dict['N'],data_dict['O'] = self.get_premium_details()
        print(self.get_insurance_details())
        print(self.get_premium_details())
        print(self.get_agency_details())
        print(self.get_bike_details())
        print(self.get_basic_info())
        print(self.get_policy_type())
        self.raw_data.close()
        return data_dict


#tb = e.get_a_table(['Total Own Damage', 'Basic Third Party Liability','Total Liability Premium'])
#data = [[cell.text for cell in row.cells] for row in tb.rows]
#print(data)

